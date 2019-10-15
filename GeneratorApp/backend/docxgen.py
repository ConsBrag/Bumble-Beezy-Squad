import docx
import os

def generateDocx(name, lastname, refer):
    author = name + " " + lastname

    os.chdir('MSword')
    doc = docx.Document('template.docx') 
    doc.paragraphs[15].text = author
    doc.add_paragraph(refer)
    doc.add_paragraph(refer)
    doc.save('change.docx')
    return 1

#Пример ввода данных
#generateDocx('Вадим','Ляпушин', "iasdjgiouahs dgiuash giausdh gasidu ghisadu ghaisud ghaipuhaihndiugahsdiogj asdhsduihasdsodihfgiausdgh")